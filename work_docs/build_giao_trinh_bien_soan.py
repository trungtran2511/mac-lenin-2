from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "D:/DemoCodeTriet/my-mln-learning-2/work_docs/GIAO_TRINH_FULL_bien_soan_6_chuong.docx"


chapters = [
    {
        "title": "Chương 1. Đối tượng, phương pháp nghiên cứu và chức năng của kinh tế chính trị Mác - Lênin",
        "goal": [
            "Hiểu quá trình hình thành và phát triển của kinh tế chính trị Mác - Lênin trong dòng chảy tư tưởng kinh tế nhân loại.",
            "Xác định được đối tượng, mục đích, phương pháp nghiên cứu và các chức năng cơ bản của môn học.",
            "Biết vận dụng tri thức kinh tế chính trị để nhìn nhận các quan hệ lợi ích, sản xuất và trao đổi trong đời sống kinh tế - xã hội.",
        ],
        "theory": [
            ("Khái quát sự hình thành", [
                "Kinh tế chính trị ra đời từ nhu cầu nhận thức các quan hệ kinh tế của xã hội. Từ chủ nghĩa trọng thương, chủ nghĩa trọng nông, kinh tế chính trị cổ điển Anh đến C.Mác, các phạm trù như hàng hóa, giá trị, tiền tệ, lợi nhuận, địa tô ngày càng được giải thích sâu hơn.",
                "C.Mác và Ph.Ăngghen kế thừa có phê phán những thành tựu trước đó, đồng thời phát hiện bản chất của phương thức sản xuất tư bản chủ nghĩa, đặc biệt thông qua học thuyết giá trị thặng dư. V.I.Lênin tiếp tục phát triển lý luận trong bối cảnh chủ nghĩa tư bản độc quyền và thời kỳ quá độ lên chủ nghĩa xã hội.",
            ]),
            ("Đối tượng nghiên cứu", [
                "Kinh tế chính trị Mác - Lênin nghiên cứu các quan hệ xã hội của sản xuất và trao đổi trong sự liên hệ biện chứng với lực lượng sản xuất và kiến trúc thượng tầng.",
                "Môn học không chỉ mô tả hiện tượng kinh tế bề ngoài mà hướng tới phát hiện các quy luật chi phối sự vận động của quan hệ kinh tế ở những trình độ phát triển lịch sử nhất định.",
            ]),
            ("Mục đích và phương pháp", [
                "Mục đích cao nhất là giúp người học hiểu quy luật kinh tế khách quan, từ đó có cơ sở khoa học để tham gia đời sống kinh tế, bảo vệ lợi ích chính đáng và góp phần phát triển xã hội.",
                "Phương pháp nền tảng là phép biện chứng duy vật, kết hợp trừu tượng hóa khoa học, phân tích - tổng hợp, logic - lịch sử, thống kê, so sánh và tổng kết thực tiễn.",
            ]),
            ("Chức năng của môn học", [
                "Chức năng nhận thức: cung cấp hệ thống tri thức về bản chất, quy luật của các quan hệ kinh tế.",
                "Chức năng thực tiễn: định hướng vận dụng quy luật kinh tế vào hoạt động sản xuất, kinh doanh, quản lý và hoạch định chính sách.",
                "Chức năng tư tưởng: góp phần hình thành thế giới quan khoa học, niềm tin vào lao động, tiến bộ xã hội và mục tiêu dân giàu, nước mạnh, dân chủ, công bằng, văn minh.",
                "Chức năng phương pháp luận: cung cấp công cụ để tiếp cận các khoa học kinh tế chuyên ngành một cách bản chất và có hệ thống.",
            ]),
        ],
        "remember": [
            "Kinh tế chính trị Mác - Lênin đặt quan hệ sản xuất và trao đổi trong mối liên hệ với lực lượng sản xuất.",
            "Không học thuộc khái niệm rời rạc; cần hiểu quy luật, điều kiện lịch sử và ý nghĩa thực tiễn của từng phạm trù.",
        ],
        "questions": [
            "Phân tích sự hình thành và phát triển của kinh tế chính trị Mác - Lênin.",
            "Trình bày đối tượng, mục đích và phương pháp nghiên cứu của kinh tế chính trị Mác - Lênin.",
            "Phân tích các chức năng của kinh tế chính trị Mác - Lênin và ý nghĩa của môn học đối với sinh viên.",
        ],
        "discussion": [
            "Vì sao khi nghiên cứu hiện tượng kinh tế cần phân biệt biểu hiện bề ngoài và bản chất bên trong?",
            "Lấy một ví dụ trong đời sống thị trường hiện nay để chỉ ra quan hệ giữa lợi ích cá nhân và lợi ích xã hội.",
        ],
    },
    {
        "title": "Chương 2. Hàng hóa, thị trường và vai trò của các chủ thể tham gia thị trường",
        "goal": [
            "Nắm được lý luận của C.Mác về sản xuất hàng hóa, hàng hóa, tiền tệ và các quy luật cơ bản của kinh tế thị trường.",
            "Hiểu vai trò của thị trường và các chủ thể như người sản xuất, người tiêu dùng, doanh nghiệp, nhà nước.",
            "Biết phân tích ưu thế, khuyết tật của thị trường và yêu cầu điều tiết trong thực tiễn.",
        ],
        "theory": [
            ("Sản xuất hàng hóa và hàng hóa", [
                "Sản xuất hàng hóa ra đời khi có phân công lao động xã hội và sự tách biệt tương đối về kinh tế giữa những người sản xuất. Sản phẩm được làm ra không chủ yếu để tự dùng mà để trao đổi, mua bán.",
                "Hàng hóa có hai thuộc tính: giá trị sử dụng và giá trị. Giá trị sử dụng đáp ứng nhu cầu; giá trị biểu hiện lao động xã hội của người sản xuất kết tinh trong hàng hóa.",
                "Lao động sản xuất hàng hóa có tính hai mặt: lao động cụ thể tạo ra giá trị sử dụng; lao động trừu tượng tạo ra giá trị. Đây là phát hiện quan trọng giúp giải thích bản chất của giá trị.",
            ]),
            ("Lượng giá trị và tiền tệ", [
                "Lượng giá trị hàng hóa được đo bằng thời gian lao động xã hội cần thiết. Năng suất lao động, cường độ lao động, mức độ phức tạp của lao động ảnh hưởng đến lượng giá trị.",
                "Tiền tệ là hàng hóa đặc biệt đóng vai trò vật ngang giá chung. Các chức năng cơ bản của tiền gồm: thước đo giá trị, phương tiện lưu thông, phương tiện thanh toán, phương tiện cất trữ và tiền tệ thế giới.",
            ]),
            ("Thị trường và kinh tế thị trường", [
                "Thị trường là tổng hòa các quan hệ kinh tế trong đó các chủ thể trao đổi hàng hóa, dịch vụ, nguồn lực theo giá cả, cung - cầu và cạnh tranh.",
                "Kinh tế thị trường có ưu thế trong phân bổ nguồn lực, kích thích sáng tạo, mở rộng lựa chọn của người tiêu dùng; đồng thời có khuyết tật như phân hóa giàu nghèo, độc quyền, khủng hoảng, ô nhiễm và thông tin bất cân xứng.",
                "Các quy luật chủ yếu gồm quy luật giá trị, cung - cầu, cạnh tranh, lưu thông tiền tệ. Các quy luật này tác động khách quan nhưng cần được nhận thức và vận dụng phù hợp.",
            ]),
            ("Vai trò các chủ thể", [
                "Người sản xuất, doanh nghiệp tổ chức cung ứng hàng hóa, dịch vụ, đổi mới công nghệ và tạo việc làm.",
                "Người tiêu dùng thông qua lựa chọn tiêu dùng tác động đến cơ cấu sản xuất và định hướng thị trường.",
                "Nhà nước xây dựng thể chế, bảo đảm cạnh tranh lành mạnh, khắc phục khuyết tật thị trường, ổn định kinh tế vĩ mô và định hướng phát triển.",
            ]),
        ],
        "remember": [
            "Hàng hóa phải được hiểu đồng thời ở hai mặt: công dụng và quan hệ xã hội ẩn sau trao đổi.",
            "Thị trường hiệu quả không đồng nghĩa với tự phát tuyệt đối; cần thể chế và quản lý phù hợp.",
        ],
        "questions": [
            "Trình bày điều kiện ra đời của sản xuất hàng hóa.",
            "Phân tích hai thuộc tính của hàng hóa và tính hai mặt của lao động sản xuất hàng hóa.",
            "Trình bày bản chất, chức năng của tiền tệ.",
            "Phân tích vai trò, chức năng, ưu thế và khuyết tật của thị trường.",
            "Phân tích vai trò của các chủ thể chính tham gia thị trường.",
        ],
        "discussion": [
            "Tại sao giá cả thị trường có thể lệch khỏi giá trị nhưng vẫn chịu sự chi phối của quy luật giá trị?",
            "Một doanh nghiệp nên làm gì để cạnh tranh lành mạnh trong nền kinh tế thị trường hiện nay?",
        ],
    },
    {
        "title": "Chương 3. Giá trị thặng dư trong nền kinh tế thị trường tư bản chủ nghĩa",
        "goal": [
            "Hiểu nguồn gốc, bản chất và các phương pháp sản xuất giá trị thặng dư.",
            "Nắm được tích lũy tư bản và các hình thức biểu hiện của giá trị thặng dư.",
            "Biết phân tích quan hệ lợi ích cơ bản trong nền kinh tế thị trường tư bản chủ nghĩa.",
        ],
        "theory": [
            ("Sự chuyển hóa tiền thành tư bản", [
                "Tiền chỉ trở thành tư bản khi được vận động nhằm mục đích thu về giá trị lớn hơn giá trị ứng trước. Công thức chung của tư bản là T - H - T', trong đó T' lớn hơn T.",
                "Mâu thuẫn của công thức chung được giải thích bằng hàng hóa sức lao động. Sức lao động có giá trị và giá trị sử dụng đặc biệt: khi được sử dụng, nó có khả năng tạo ra giá trị mới lớn hơn giá trị bản thân nó.",
            ]),
            ("Nguồn gốc và bản chất giá trị thặng dư", [
                "Giá trị thặng dư là phần giá trị mới do lao động của công nhân tạo ra vượt quá giá trị sức lao động và bị nhà tư bản chiếm hữu.",
                "Tư bản bất biến chuyển giá trị cũ vào sản phẩm; tư bản khả biến thông qua sức lao động tạo ra giá trị mới và giá trị thặng dư.",
            ]),
            ("Các phương pháp sản xuất giá trị thặng dư", [
                "Giá trị thặng dư tuyệt đối được tạo ra bằng cách kéo dài ngày lao động hoặc tăng cường độ lao động khi thời gian lao động tất yếu không đổi.",
                "Giá trị thặng dư tương đối được tạo ra bằng cách rút ngắn thời gian lao động tất yếu thông qua tăng năng suất lao động xã hội, làm giảm giá trị sức lao động.",
                "Giá trị thặng dư siêu ngạch xuất hiện khi một nhà tư bản có năng suất cá biệt cao hơn mức xã hội, bán theo giá trị xã hội và thu phần chênh lệch.",
            ]),
            ("Tích lũy tư bản và biểu hiện của giá trị thặng dư", [
                "Tích lũy tư bản là biến một phần giá trị thặng dư thành tư bản phụ thêm. Quá trình này làm tăng quy mô sản xuất, thúc đẩy tích tụ và tập trung tư bản.",
                "Giá trị thặng dư biểu hiện dưới các hình thức như lợi nhuận, lợi tức, địa tô. Các hình thức này che giấu nguồn gốc trực tiếp của giá trị thặng dư là lao động không được trả công.",
            ]),
        ],
        "remember": [
            "Chìa khóa của chương là hàng hóa sức lao động và sự phân biệt lao động tất yếu - lao động thặng dư.",
            "Lợi nhuận, lợi tức, địa tô là hình thức biểu hiện; giá trị thặng dư là bản chất cần phân tích.",
        ],
        "questions": [
            "Phân tích công thức chung của tư bản và mâu thuẫn của công thức chung.",
            "Trình bày hàng hóa sức lao động và vai trò của nó trong việc tạo ra giá trị thặng dư.",
            "Phân tích nguồn gốc, bản chất của giá trị thặng dư.",
            "So sánh giá trị thặng dư tuyệt đối, tương đối và siêu ngạch.",
            "Trình bày bản chất tích lũy tư bản và một số hệ quả của tích lũy tư bản.",
            "Phân tích các hình thức biểu hiện của giá trị thặng dư: lợi nhuận, lợi tức, địa tô.",
        ],
        "discussion": [
            "Trong bối cảnh kinh tế số, năng suất lao động và giá trị thặng dư có thể được nhìn nhận như thế nào?",
            "Vì sao cần phân biệt lợi nhuận nhìn thấy trong kinh doanh với nguồn gốc xã hội của giá trị thặng dư?",
        ],
    },
    {
        "title": "Chương 4. Cạnh tranh và độc quyền trong nền kinh tế thị trường tư bản chủ nghĩa",
        "goal": [
            "Hiểu quan hệ giữa cạnh tranh, tích tụ tư bản và sự hình thành độc quyền.",
            "Nắm được lý luận của V.I.Lênin về đặc điểm kinh tế của độc quyền và độc quyền nhà nước.",
            "Phân tích biểu hiện mới của độc quyền trong điều kiện hiện nay và vai trò lịch sử của chủ nghĩa tư bản.",
        ],
        "theory": [
            ("Cạnh tranh và sự hình thành độc quyền", [
                "Cạnh tranh là quan hệ ganh đua giữa các chủ thể kinh tế nhằm giành điều kiện sản xuất, tiêu thụ và lợi nhuận thuận lợi hơn. Cạnh tranh vừa thúc đẩy đổi mới vừa dẫn tới phân hóa doanh nghiệp.",
                "Tích tụ và tập trung sản xuất làm xuất hiện các tổ chức độc quyền. Khi một số ít doanh nghiệp lớn chi phối thị trường, cạnh tranh tự do chuyển sang cạnh tranh trong điều kiện độc quyền.",
            ]),
            ("Độc quyền và tác động của độc quyền", [
                "Độc quyền có thể tạo điều kiện tập trung nguồn lực cho đổi mới kỹ thuật, mở rộng quy mô và giảm chi phí. Tuy nhiên, độc quyền cũng có thể hạn chế cạnh tranh, áp đặt giá cả, kìm hãm sáng tạo và gây thiệt hại cho người tiêu dùng.",
                "Vì vậy, kiểm soát độc quyền, chống lạm dụng vị trí thống lĩnh và bảo vệ cạnh tranh là yêu cầu quan trọng của quản lý kinh tế hiện đại.",
            ]),
            ("Lý luận của V.I.Lênin", [
                "V.I.Lênin chỉ ra những đặc điểm kinh tế cơ bản của chủ nghĩa tư bản độc quyền: tập trung sản xuất và tư bản dẫn tới độc quyền; tư bản tài chính và đầu sỏ tài chính; xuất khẩu tư bản; liên minh độc quyền quốc tế; phân chia thế giới về kinh tế và lãnh thổ.",
                "Độc quyền nhà nước là sự kết hợp sức mạnh của các tổ chức độc quyền với sức mạnh của nhà nước tư sản nhằm phục vụ lợi ích của tư bản độc quyền trong những điều kiện lịch sử nhất định.",
            ]),
            ("Biểu hiện mới hiện nay", [
                "Độc quyền hiện đại gắn với công nghệ, dữ liệu, sở hữu trí tuệ, nền tảng số, chuỗi giá trị toàn cầu và quyền lực thị trường xuyên quốc gia.",
                "Chủ nghĩa tư bản có vai trò lịch sử trong phát triển lực lượng sản xuất, xã hội hóa sản xuất và thúc đẩy khoa học - công nghệ, nhưng đồng thời chứa đựng mâu thuẫn về bất bình đẳng, khủng hoảng, độc quyền và giới hạn phát triển bền vững.",
            ]),
        ],
        "remember": [
            "Độc quyền không xóa bỏ cạnh tranh; nó làm cạnh tranh diễn ra trong hình thức và điều kiện mới.",
            "Cần nhìn độc quyền ở cả hai mặt: khả năng tổ chức quy mô lớn và nguy cơ lạm dụng quyền lực thị trường.",
        ],
        "questions": [
            "Phân tích nguyên nhân hình thành độc quyền trong nền kinh tế thị trường tư bản chủ nghĩa.",
            "Trình bày tác động tích cực và tiêu cực của độc quyền đối với nền kinh tế.",
            "Vì sao cần kiểm soát độc quyền? Có thể kiểm soát bằng những phương thức nào?",
            "Phân tích các đặc điểm kinh tế của chủ nghĩa tư bản độc quyền theo V.I.Lênin.",
            "Phân tích độc quyền nhà nước trong chủ nghĩa tư bản và biểu hiện mới hiện nay.",
            "Trình bày vai trò lịch sử và giới hạn của chủ nghĩa tư bản.",
        ],
        "discussion": [
            "Các nền tảng số lớn hiện nay có những dấu hiệu độc quyền nào?",
            "Chính sách cạnh tranh nên cân bằng thế nào giữa khuyến khích đổi mới và chống lạm dụng vị trí thống lĩnh?",
        ],
    },
    {
        "title": "Chương 5. Kinh tế thị trường định hướng xã hội chủ nghĩa và các quan hệ lợi ích kinh tế ở Việt Nam",
        "goal": [
            "Hiểu tính tất yếu, khái niệm và đặc trưng của kinh tế thị trường định hướng xã hội chủ nghĩa ở Việt Nam.",
            "Nắm nhiệm vụ hoàn thiện thể chế kinh tế thị trường định hướng xã hội chủ nghĩa.",
            "Phân tích quan hệ lợi ích kinh tế và vai trò của nhà nước trong bảo đảm hài hòa lợi ích.",
        ],
        "theory": [
            ("Tính tất yếu và khái niệm", [
                "Phát triển kinh tế thị trường định hướng xã hội chủ nghĩa ở Việt Nam là kết quả của đổi mới tư duy kinh tế, phù hợp với yêu cầu phát triển lực lượng sản xuất, hội nhập quốc tế và mục tiêu xây dựng chủ nghĩa xã hội.",
                "Đây là nền kinh tế vận hành đầy đủ, đồng bộ theo các quy luật của kinh tế thị trường, đồng thời được định hướng bởi mục tiêu dân giàu, nước mạnh, dân chủ, công bằng, văn minh dưới sự quản lý của nhà nước pháp quyền xã hội chủ nghĩa và sự lãnh đạo của Đảng.",
            ]),
            ("Đặc trưng cơ bản", [
                "Nhiều hình thức sở hữu, nhiều thành phần kinh tế cùng phát triển; kinh tế nhà nước giữ vai trò chủ đạo trong những lĩnh vực then chốt.",
                "Thị trường giữ vai trò quan trọng trong phân bổ nguồn lực, còn nhà nước định hướng, điều tiết, khắc phục khuyết tật thị trường và bảo đảm tiến bộ, công bằng xã hội.",
                "Tăng trưởng kinh tế gắn với phát triển văn hóa, con người, bảo vệ môi trường và bảo đảm an sinh xã hội.",
            ]),
            ("Hoàn thiện thể chế", [
                "Hoàn thiện pháp luật về sở hữu, doanh nghiệp, cạnh tranh, thị trường các yếu tố sản xuất, tài chính, đất đai, khoa học - công nghệ và lao động.",
                "Nâng cao hiệu lực quản lý nhà nước; phát triển đồng bộ các loại thị trường; cải thiện môi trường kinh doanh; bảo đảm minh bạch, kỷ luật, kỷ cương và trách nhiệm giải trình.",
            ]),
            ("Quan hệ lợi ích kinh tế", [
                "Lợi ích kinh tế là động lực trực tiếp thúc đẩy hoạt động của các chủ thể. Các quan hệ lợi ích chủ yếu gồm lợi ích cá nhân, tập thể, doanh nghiệp, xã hội, nhà nước, người lao động và người sử dụng lao động.",
                "Các lợi ích vừa thống nhất vừa có thể mâu thuẫn. Nhà nước có vai trò xây dựng thể chế, phân phối và điều tiết, bảo vệ lợi ích hợp pháp, xử lý xung đột và hướng lợi ích riêng vào lợi ích chung.",
            ]),
        ],
        "remember": [
            "Định hướng xã hội chủ nghĩa không phủ nhận thị trường; nó quy định mục tiêu, nguyên tắc và cách nhà nước điều tiết thị trường.",
            "Hài hòa lợi ích là điều kiện để phát triển bền vững và giữ ổn định xã hội.",
        ],
        "questions": [
            "Phân tích tính tất yếu khách quan của việc phát triển kinh tế thị trường định hướng xã hội chủ nghĩa ở Việt Nam.",
            "Trình bày những đặc trưng của kinh tế thị trường định hướng xã hội chủ nghĩa ở Việt Nam.",
            "Phân tích các nhiệm vụ chủ yếu để hoàn thiện thể chế kinh tế thị trường định hướng xã hội chủ nghĩa ở Việt Nam.",
            "Trình bày khái niệm, đặc trưng và các nhân tố ảnh hưởng đến quan hệ lợi ích kinh tế.",
            "Phân tích các quan hệ lợi ích kinh tế chủ yếu trong nền kinh tế thị trường.",
            "Phân tích vai trò của nhà nước trong bảo đảm hài hòa các lợi ích kinh tế.",
        ],
        "discussion": [
            "Với tư cách công dân, sinh viên cần làm gì để góp phần hoàn thiện thể chế kinh tế thị trường định hướng xã hội chủ nghĩa?",
            "Khi lợi ích cá nhân và lợi ích cộng đồng có xung đột, cần xử lý theo nguyên tắc nào?",
        ],
    },
    {
        "title": "Chương 6. Công nghiệp hóa, hiện đại hóa và hội nhập kinh tế quốc tế của Việt Nam",
        "goal": [
            "Hiểu nội dung cơ bản của công nghiệp hóa, hiện đại hóa ở Việt Nam.",
            "Nắm yêu cầu thích ứng với cuộc Cách mạng công nghiệp lần thứ tư.",
            "Phân tích tính tất yếu, tác động và giải pháp nâng cao hiệu quả hội nhập kinh tế quốc tế.",
        ],
        "theory": [
            ("Công nghiệp hóa, hiện đại hóa", [
                "Công nghiệp hóa, hiện đại hóa là quá trình chuyển đổi căn bản, toàn diện hoạt động sản xuất, kinh doanh, dịch vụ và quản lý từ sử dụng lao động thủ công là chính sang sử dụng phổ biến công nghệ, phương tiện và phương pháp tiên tiến.",
                "Ở Việt Nam, công nghiệp hóa, hiện đại hóa gắn với phát triển kinh tế tri thức, đổi mới mô hình tăng trưởng, nâng cao năng suất lao động, chất lượng nguồn nhân lực và năng lực cạnh tranh quốc gia.",
            ]),
            ("Cách mạng công nghiệp lần thứ tư", [
                "Cách mạng công nghiệp 4.0 dựa trên nền tảng số, trí tuệ nhân tạo, dữ liệu lớn, Internet vạn vật, tự động hóa, công nghệ sinh học và vật liệu mới.",
                "Cơ hội là tăng năng suất, mở rộng mô hình kinh doanh mới, nâng cao chất lượng quản trị và kết nối toàn cầu. Thách thức là nguy cơ tụt hậu, mất việc làm giản đơn, khoảng cách số, yêu cầu kỹ năng mới và an ninh dữ liệu.",
            ]),
            ("Hội nhập kinh tế quốc tế", [
                "Hội nhập kinh tế quốc tế là quá trình một quốc gia gắn kết nền kinh tế của mình với kinh tế khu vực và thế giới trên cơ sở chia sẻ lợi ích, tuân thủ các chuẩn mực và cam kết chung.",
                "Hội nhập tạo cơ hội mở rộng thị trường, thu hút vốn, công nghệ, kinh nghiệm quản lý, nâng cao năng lực cạnh tranh; đồng thời làm gia tăng cạnh tranh, phụ thuộc bên ngoài và rủi ro phân hóa lợi ích.",
            ]),
            ("Giải pháp phát triển", [
                "Hoàn thiện thể chế, nâng cao chất lượng chiến lược và lộ trình hội nhập; phát triển nguồn nhân lực chất lượng cao; đầu tư khoa học - công nghệ, hạ tầng số và đổi mới sáng tạo.",
                "Chủ động tham gia các liên kết kinh tế quốc tế, nâng cao năng lực cạnh tranh của doanh nghiệp, xây dựng nền kinh tế độc lập, tự chủ gắn với hội nhập sâu rộng và hiệu quả.",
            ]),
        ],
        "remember": [
            "Công nghiệp hóa, hiện đại hóa hiện nay phải gắn với kinh tế tri thức, chuyển đổi số và phát triển bền vững.",
            "Hội nhập hiệu quả cần đi cùng độc lập, tự chủ; mở cửa nhưng phải nâng được năng lực nội sinh.",
        ],
        "questions": [
            "Phân tích nội dung cơ bản của quá trình công nghiệp hóa, hiện đại hóa ở Việt Nam.",
            "Phân tích quan điểm và giải pháp thực hiện công nghiệp hóa, hiện đại hóa trong bối cảnh Cách mạng công nghiệp lần thứ tư.",
            "Phân tích tính tất yếu của hội nhập kinh tế quốc tế và những tác động của hội nhập đối với Việt Nam.",
            "Trình bày giải pháp nhằm nâng cao hiệu quả hội nhập kinh tế quốc tế trong phát triển của Việt Nam.",
        ],
        "discussion": [
            "Tác động tích cực và tiêu cực của hội nhập kinh tế quốc tế đối với sinh viên và người lao động trẻ là gì?",
            "Việt Nam cần làm gì để vừa hội nhập sâu rộng vừa giữ vững nền kinh tế độc lập, tự chủ?",
        ],
    },
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_font(run, name="Calibri", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_bullet(doc, text, style="List Bullet"):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_font(r)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_font(r)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    set_font(r)
    return p


def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_font(run, size={1: 16, 2: 13, 3: 12}.get(level, 11), bold=True, color="2E74B5" if level < 3 else "1F4D78")
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_callout(doc, title, items):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_width(cell, 9360)
    set_cell_shading(cell, "F4F6F9")
    p = cell.paragraphs[0]
    r = p.add_run(title)
    set_font(r, bold=True, color="1F4D78")
    for item in items:
        p = cell.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(item)
        set_font(r, size=10.5)
    doc.add_paragraph()


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    styles["Normal"]._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    styles["Normal"].font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    r = title.add_run("GIÁO TRÌNH KINH TẾ CHÍNH TRỊ MÁC - LÊNIN")
    set_font(r, size=18, bold=True, color="0B2545")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Bản biên soạn lại theo 6 chương: bổ sung phần lý luận, tách câu hỏi ôn tập và chia nhỏ nội dung để dễ học")
    set_font(r, size=12, color="555555")

    add_callout(doc, "Cách dùng tài liệu", [
        "Đọc mục tiêu trước khi học chương.",
        "Nắm phần lý luận cốt lõi, sau đó dùng mục ghi nhớ để tự kiểm tra.",
        "Làm câu hỏi ôn tập để củng cố kiến thức và dùng câu hỏi thảo luận cho học nhóm hoặc chuẩn bị bài.",
    ])

    add_heading(doc, "Mục lục nội dung", 1)
    for ch in chapters:
        add_bullet(doc, ch["title"])

    for idx, ch in enumerate(chapters, 1):
        if idx > 1:
            doc.add_section(WD_SECTION_START.NEW_PAGE)
        add_heading(doc, ch["title"], 1)

        add_heading(doc, "1. Mục tiêu học tập", 2)
        for item in ch["goal"]:
            add_bullet(doc, item)

        add_heading(doc, "2. Phần lý luận cần nắm", 2)
        for heading, paragraphs in ch["theory"]:
            add_heading(doc, heading, 3)
            for para in paragraphs:
                add_body(doc, para)

        add_heading(doc, "3. Ghi nhớ nhanh", 2)
        add_callout(doc, "Ý chính của chương", ch["remember"])

        add_heading(doc, "4. Câu hỏi ôn tập", 2)
        for q in ch["questions"]:
            add_number(doc, q)

        add_heading(doc, "5. Câu hỏi thảo luận/vận dụng", 2)
        for q in ch["discussion"]:
            add_bullet(doc, q)

    add_heading(doc, "Phụ lục. Gợi ý cách học theo 6 chương", 1)
    add_body(doc, "Người học nên đi theo trình tự: chương 1 để nắm phương pháp luận; chương 2 để hiểu hàng hóa, thị trường và chủ thể; chương 3 để phân tích giá trị thặng dư; chương 4 để hiểu cạnh tranh, độc quyền; chương 5 để liên hệ mô hình kinh tế Việt Nam; chương 6 để vận dụng vào công nghiệp hóa, hiện đại hóa và hội nhập.")
    add_callout(doc, "Mẫu ôn tập mỗi chương", [
        "Tóm tắt chương trong 5 ý chính.",
        "Viết lại 3 khái niệm quan trọng bằng lời của mình.",
        "Chọn 1 câu hỏi ôn tập để lập dàn ý.",
        "Liên hệ 1 ví dụ thực tế ở Việt Nam hoặc trong đời sống sinh viên.",
    ])

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.text = "Bản biên soạn học tập từ giáo trình gốc - chia 6 chương"
        for run in footer.runs:
            set_font(run, size=9, color="666666")

    doc.save(OUT)


if __name__ == "__main__":
    build()
